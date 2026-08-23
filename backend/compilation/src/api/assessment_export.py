from io import BytesIO

from docx import Document
import httpx
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from xml.sax.saxutils import escape

from src.core.dependency import get_current_user
from src.core.infrastructure.configuration import settings


router = APIRouter(prefix="/exports/assessment", tags=["assessment-export"])


def text_projection(node):
    values = []
    if isinstance(node, dict):
        if node.get("type") == "text" and isinstance(node.get("text"), str):
            values.append(node["text"])
        if node.get("type") in {"inlineMath", "blockMath", "mathematics"}:
            latex = node.get("attrs", {}).get("latex") or node.get("attrs", {}).get("content")
            if latex:
                values.append(str(latex))
        if node.get("type") == "image" and node.get("attrs", {}).get("alt"):
            values.append(f"[{node['attrs']['alt']}]")
        for key, value in node.items():
            if key not in {"text", "attrs"}:
                values.extend(text_projection(value))
    elif isinstance(node, list):
        for value in node:
            values.extend(text_projection(value))
    return values


def item_text(question):
    stem = " ".join(text_projection(question.get("stem_doc", {})))
    options = [
        f"{option['id']} {' '.join(text_projection(option.get('content_doc', {})))}"
        for option in question.get("options", [])
    ]
    return stem, options


async def load_snapshot(version_id, current_user):
    async with httpx.AsyncClient(timeout=20) as client:
        response = await client.get(
            f"{settings.ASSESSMENT_URL}/assessment-versions/{version_id}/internal/export-snapshot",
            headers={
                "X-Internal-Token": settings.SECRET_KEY,
                "X-Actor-Id": current_user.id,
                "X-Actor-Role": current_user.role.value,
            },
        )
    if response.status_code in {403, 404, 409}:
        raise HTTPException(status_code=response.status_code, detail=response.json().get("detail"))
    if response.status_code >= 400:
        raise HTTPException(status_code=503, detail={"code": "assessment_snapshot_unavailable"})
    snapshot = response.json()
    version = snapshot["version"]
    questions = snapshot["questions"]
    return version, {question["_id"]: question for question in questions}


@router.post("/{version_id}/pdf")
async def export_assessment_pdf(version_id: str, current_user=Depends(get_current_user)):
    version, questions = await load_snapshot(version_id, current_user)
    output = BytesIO()
    pdfmetrics.registerFont(TTFont("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("AssessmentTitle", parent=styles["Title"], fontName="DejaVu")
    heading_style = ParagraphStyle("QuestionHeading", parent=styles["Heading2"], fontName="DejaVu")
    body_style = ParagraphStyle(
        "QuestionBody", parent=styles["BodyText"], fontName="DejaVu", leading=16
    )
    story = [Paragraph(escape(version["title"]), title_style), Spacer(1, 6 * mm)]
    for item in version["items"]:
        stem, options = item_text(questions[item["question_version_id"]])
        story.extend(
            [
                Paragraph(f"Câu {item['position']}", heading_style),
                Paragraph(escape(stem), body_style),
            ]
        )
        story.extend(Paragraph(escape(option), body_style) for option in options)
        story.append(Spacer(1, 4 * mm))
    SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    ).build(story)
    return Response(
        output.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{version_id}.pdf"'},
    )


@router.post("/{version_id}/docx")
async def export_assessment_docx(version_id: str, current_user=Depends(get_current_user)):
    version, questions = await load_snapshot(version_id, current_user)
    document = Document()
    document.add_heading(version["title"], level=1)
    for item in version["items"]:
        stem, options = item_text(questions[item["question_version_id"]])
        document.add_heading(f"Câu {item['position']}", level=2)
        document.add_paragraph(stem)
        for option in options:
            document.add_paragraph(option)
    output = BytesIO()
    document.save(output)
    return Response(
        output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{version_id}.docx"'},
    )
